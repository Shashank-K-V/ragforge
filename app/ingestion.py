"""
ingestion.py — Document loading, cleaning, and chunking logic.

Responsibilities
----------------
1. Accept an uploaded file (PDF, TXT, DOCX) and persist it to UPLOAD_DIR.
2. Extract raw text, preserving page metadata for PDFs.
3. Split text into overlapping chunks using LangChain's RecursiveCharacterTextSplitter
   (character-based, which maps roughly to token counts for English text).
4. Return LangChain Document objects ready for embedding in retrieval.py.

Design decisions
----------------
* RecursiveCharacterTextSplitter is preferred over TokenTextSplitter because
  it doesn't require a tokeniser dependency, is deterministic, and respects
  paragraph / sentence boundaries before resorting to hard character cuts.
* CHUNK_SIZE = 2000 chars ≈ 500 tokens (at ~4 chars/token for English).
* CHUNK_OVERLAP = 200 chars ≈ 50 tokens — large enough to preserve a
  complete sentence at each boundary without wasting too much space.
* Each chunk carries metadata: document_id, filename, chunk_index, page_number
  (PDF only).  These travel with the embedding into ChromaDB and are
  returned in QueryResponse.source_documents.
"""

from __future__ import annotations

import io
import logging
import uuid
from pathlib import Path
from typing import IO

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.config import settings
from app.models import DocumentType

logger = logging.getLogger(__name__)


# ======================================================================= #
#  Helpers                                                                  #
# ======================================================================= #


def _detect_document_type(filename: str) -> DocumentType:
    """Infer DocumentType from the file extension (lower-cased)."""
    ext = Path(filename).suffix.lower().lstrip(".")
    mapping = {
        "pdf": DocumentType.PDF,
        "txt": DocumentType.TXT,
        "text": DocumentType.TXT,
        "docx": DocumentType.DOCX,
    }
    return mapping.get(ext, DocumentType.UNKNOWN)


def _ensure_upload_dir() -> Path:
    """Create UPLOAD_DIR if it doesn't exist and return its Path."""
    upload_path = Path(settings.UPLOAD_DIR)
    upload_path.mkdir(parents=True, exist_ok=True)
    return upload_path


# ======================================================================= #
#  Per-format text extractors                                               #
# ======================================================================= #


def _extract_text_from_pdf(file_bytes: bytes) -> list[tuple[str, int]]:
    """
    Extract text from a PDF, returning a list of (page_text, page_number) tuples.

    Uses PyPDF2.  Falls back gracefully if a page can't be decoded.
    """
    try:
        import PyPDF2  # imported lazily so the app still starts without it
    except ImportError as exc:
        raise RuntimeError(
            "PyPDF2 is required for PDF ingestion. Install it with: pip install PyPDF2"
        ) from exc

    pages: list[tuple[str, int]] = []
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))

    for page_num, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            logger.warning("Could not extract text from PDF page %d: %s", page_num, exc)
            text = ""
        if text.strip():
            pages.append((text, page_num))

    return pages


def _extract_text_from_docx(file_bytes: bytes) -> list[tuple[str, int]]:
    """
    Extract text from a DOCX file.

    python-docx doesn't expose page numbers, so we assign page_number = 0
    (meaning "unknown / not applicable") for every paragraph block.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX ingestion. "
            "Install it with: pip install python-docx"
        ) from exc

    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [(full_text, 0)]


def _extract_text_from_txt(file_bytes: bytes) -> list[tuple[str, int]]:
    """Decode a plain-text file, trying UTF-8 then Latin-1 as fallback."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")
    return [(text, 0)]


# ======================================================================= #
#  Public ingestion API                                                     #
# ======================================================================= #


def save_upload(filename: str, file_obj: IO[bytes]) -> tuple[str, Path]:
    """
    Persist an uploaded file to UPLOAD_DIR under a UUID sub-directory.

    Returns
    -------
    document_id : str
        Stable UUID for this document, used as the ChromaDB collection key.
    saved_path : Path
        Absolute path to the saved file on disk.
    """
    document_id = str(uuid.uuid4())
    upload_dir = _ensure_upload_dir() / document_id
    upload_dir.mkdir(parents=True, exist_ok=True)

    saved_path = upload_dir / filename
    content = file_obj.read()

    # Enforce max file size
    max_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
    if len(content) > max_bytes:
        raise ValueError(
            f"File size {len(content) / 1_048_576:.1f} MB exceeds the "
            f"{settings.MAX_FILE_SIZE_MB} MB limit."
        )

    saved_path.write_bytes(content)
    logger.info("Saved upload: %s → %s", filename, saved_path)
    return document_id, saved_path


def load_and_chunk(
    document_id: str,
    filename: str,
    file_bytes: bytes,
    *,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """
    Load raw bytes, extract text, and split into overlapping chunks.

    Parameters
    ----------
    document_id:
        UUID assigned at upload time; embedded in every chunk's metadata.
    filename:
        Original file name; used to detect format and stored in metadata.
    file_bytes:
        Raw bytes of the uploaded file.
    chunk_size:
        Character length per chunk.  Defaults to settings.CHUNK_SIZE.
    chunk_overlap:
        Character overlap between consecutive chunks.
        Defaults to settings.CHUNK_OVERLAP.

    Returns
    -------
    list[Document]
        LangChain Document objects, each carrying:
          - page_content : the chunk text
          - metadata     : {document_id, filename, chunk_index, page_number}
    """
    chunk_size = chunk_size or settings.CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
    doc_type = _detect_document_type(filename)

    # --- Extract (page_text, page_number) pairs ---
    if doc_type == DocumentType.PDF:
        pages = _extract_text_from_pdf(file_bytes)
    elif doc_type == DocumentType.DOCX:
        pages = _extract_text_from_docx(file_bytes)
    else:
        # TXT and UNKNOWN both fall through to plain-text extraction
        pages = _extract_text_from_txt(file_bytes)

    if not pages or not any(text.strip() for text, _ in pages):
        raise ValueError(f"No extractable text found in '{filename}'.")

    # --- Build LangChain Documents (one per page before chunking) ---
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        # Try to break at paragraph → sentence → word boundaries first
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )

    all_chunks: list[Document] = []
    chunk_index = 0

    for page_text, page_number in pages:
        page_doc = Document(
            page_content=page_text,
            metadata={
                "document_id": document_id,
                "filename": filename,
                "page_number": page_number,
            },
        )
        # split_documents returns new Document objects, inheriting metadata
        page_chunks = splitter.split_documents([page_doc])

        for chunk in page_chunks:
            chunk.metadata["chunk_index"] = chunk_index
            all_chunks.append(chunk)
            chunk_index += 1

    logger.info(
        "Ingested '%s' (type=%s): %d page(s) → %d chunk(s)",
        filename,
        doc_type.value,
        len(pages),
        len(all_chunks),
    )
    return all_chunks


def ingest_file(
    filename: str,
    file_bytes: bytes,
    document_id: str | None = None,
) -> tuple[str, list[Document]]:
    """
    High-level convenience wrapper used by main.py.

    Parameters
    ----------
    filename    : Original file name (used for type detection + metadata).
    file_bytes  : Raw file content.
    document_id : Pre-assigned ID; generates a fresh UUID if None.

    Returns
    -------
    (document_id, chunks) where chunks are ready for embedding.
    """
    if document_id is None:
        document_id = str(uuid.uuid4())

    chunks = load_and_chunk(document_id, filename, file_bytes)
    return document_id, chunks
